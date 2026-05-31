"""Generate generic Stage 2 Word docs from sanitized JSON content.

This public template intentionally does not contain real project paths or
project-specific lesson text. Prepare a UTF-8 JSON file and pass it with
--content. The AI agent can create that JSON from the current courseware.
"""
import argparse
import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BLACK = RGBColor(0, 0, 0)


def set_font(run, name="微软雅黑", size=Pt(11), bold=False, color=BLACK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    sizes = {0: Pt(18), 1: Pt(15), 2: Pt(13), 3: Pt(12)}
    set_font(run, size=sizes.get(level, Pt(11)), bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.space_before = Pt(12)
    p.space_after = Pt(6)
    return p


def add_body(doc, text="", bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=Pt(11), bold=bold)
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    return p


def setup_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    return doc


def build_script_doc(data):
    title = data["title"]
    doc = setup_doc()
    add_heading(doc, f"逐字稿——《{title}》", 0)

    for line in data.get("meta_lines", []):
        add_body(doc, line)
    add_body(doc)

    for page in data.get("script_pages", []):
        page_no = page.get("page_no", "")
        page_title = page.get("title", "")
        add_heading(doc, f"第{page_no}页｜{page_title}" if page_no else page_title, 1)
        for key, label in [
            ("purpose", "页面目的"),
            ("teacher_talk", "教师话术"),
            ("questions", "课堂提问"),
            ("answers", "学生可能回答"),
            ("activity", "活动/操作说明"),
            ("transition", "转场语"),
        ]:
            value = page.get(key)
            if value:
                add_body(doc, f"{label}：{value}")
        add_body(doc)
    return doc


def build_design_doc(data):
    title = data["title"]
    doc = setup_doc()
    add_heading(doc, f"教学设计——《{title}》", 0)

    for section in data.get("design_sections", []):
        add_heading(doc, section.get("heading", ""), section.get("level", 1))
        for item in section.get("items", []):
            add_body(doc, item)
        add_body(doc)
    return doc


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--content", required=True, help="UTF-8 JSON content file")
    parser.add_argument("--out-dir", required=True, help="Output folder")
    parser.add_argument("--base-name", required=True, help="Base name used in output filenames")
    args = parser.parse_args()

    with open(args.content, "r", encoding="utf-8") as f:
        data = json.load(f)

    os.makedirs(args.out_dir, exist_ok=True)
    script_path = os.path.join(args.out_dir, f"2.逐字稿-{args.base_name}.docx")
    design_path = os.path.join(args.out_dir, f"3.教学设计-{args.base_name}.docx")

    build_script_doc(data).save(script_path)
    build_design_doc(data).save(design_path)

    print(f"Saved: {script_path}")
    print(f"Saved: {design_path}")


if __name__ == "__main__":
    main()
